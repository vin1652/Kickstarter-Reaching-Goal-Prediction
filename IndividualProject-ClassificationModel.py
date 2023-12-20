#############################################################################################
## Developing the Classification model ############
#############################################################################################
# Load Libraries
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.preprocessing import StandardScaler
import pandas as pd
import numpy as np
from sklearn.metrics import accuracy_score
from sklearn.ensemble import RandomForestClassifier
from sklearn.ensemble import IsolationForest
import matplotlib.pyplot as plt

# Import data
kickstarter_df = pd.read_excel("C:/Users/vinay/OneDrive/Documents/.spyder-py3/Kickstarter.xlsx")

# Pre-Processing
kickstarter_df = kickstarter_df.dropna()

# Filtering the DataFrame to keep only rows where 'status' is 'successful' or 'failed'
kickstarter_df = kickstarter_df[kickstarter_df['state'].isin(['successful', 'failed'])]

kickstarter_df['goal_fixed']=kickstarter_df['goal']*kickstarter_df['static_usd_rate']
kickstarter_df = kickstarter_df.drop(columns=['id', 'name', 'pledged', 'state_changed_at', 'backers_count', 'usd_pledged',
                                              'state_changed_at_weekday', 'state_changed_at_month', 'state_changed_at_day', 
                                              'state_changed_at_yr', 'state_changed_at_hr', 'launch_to_state_change_days',
                                              'deadline', 'created_at', 'launched_at','goal','spotlight','staff_pick'])
y = kickstarter_df['state']
kickstarter_df = kickstarter_df.drop('state', axis=1)

categorical_features = kickstarter_df.select_dtypes(include=['object']).columns
X = pd.get_dummies(kickstarter_df, columns=categorical_features, drop_first=True)

# Splitting the dataset into the Training set and Test set
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=5)

# Using Isolation Forest for detecting outliers
iso_forest = IsolationForest(n_estimators=100, contamination='auto', random_state=5)  # Let the algorithm determine the contamination
iso_forest.fit(X_train)

# Predict if a data point is an outlier. -1 for outliers, 1 for inliers.
outlier_predictions = iso_forest.predict(X_train)

# Select all rows that are not outliers
is_inlier = outlier_predictions == 1
X_train_filtered = X_train[is_inlier]
y_train_filtered = y_train[is_inlier]

# RandomForest with GridSearchCV
param_grid = {
    'n_estimators': [50,100, 150, 200, 250, 300],
    'max_features': ['auto', 'sqrt', 0.5, 0.75,3,4,5,6] ,
    'max_depth':[10, 20, 30, None],
    'min_samples_split': [2, 4, 6, 8,10],
    'min_samples_leaf': [1, 2,3, 4, 6]
}

rf = RandomForestClassifier(random_state=2000)
grid_search = GridSearchCV(estimator=rf, param_grid=param_grid, cv=5, n_jobs=-1, verbose=2, scoring='accuracy')
grid_search.fit(X_train_filtered, y_train_filtered)

# Best parameters and model
best_params = grid_search.best_params_
best_rf = grid_search.best_estimator_

# Evaluating the best model
predictions = best_rf.predict(X_test)
accuracy = accuracy_score(y_test, predictions)

print("Best Parameters:", best_params)
print("Best Model Accuracy:", accuracy)

features = X.columns
importance_df = pd.DataFrame({'Feature': features, 'Importance': best_rf.feature_importances_})
importance_df = importance_df.sort_values(by='Importance', ascending=False).reset_index(drop=True)

# Sort the DataFrame by importance and select top 10
top_10_importance_df = importance_df.sort_values(by='Importance', ascending=False).head(10)

# Plotting
plt.figure(figsize=(10, 6))
plt.barh(top_10_importance_df['Feature'], top_10_importance_df['Importance'], color='skyblue')
plt.xlabel('Importance')
plt.ylabel('Feature')
plt.title('Top 10 Feature Importances in Random Forest Model')
plt.gca().invert_yaxis()  # To display the highest importance at the top
plt.show()

########################################################################################
## Code for Grading #################
#########################################################################################
# Import Grading Data
kickstarter_grading_df = pd.read_excel("C:/Users/vinay/OneDrive/Documents/.spyder-py3/Kickstarter-Grading-Sample.xlsx")#"Kickstarter-Grading-Sample.xlsx")

# Pre-Processing
kickstarter_grading_df = kickstarter_grading_df.dropna()

# Filtering the DataFrame to keep only rows where 'status' is 'successful' or 'failed'
kickstarter_grading_df = kickstarter_grading_df[kickstarter_grading_df['state'].isin(['successful', 'failed'])]

kickstarter_grading_df['goal_fixed']=kickstarter_grading_df['goal']*kickstarter_grading_df['static_usd_rate']
kickstarter_grading_df = kickstarter_grading_df.drop(columns=['id', 'name', 'pledged', 'state_changed_at', 'backers_count', 'usd_pledged',
                                              'state_changed_at_weekday', 'state_changed_at_month', 'state_changed_at_day', 
                                              'state_changed_at_yr', 'state_changed_at_hr', 'launch_to_state_change_days',
                                              'deadline', 'created_at', 'launched_at','goal','spotlight','staff_pick'])
y_grading = kickstarter_grading_df['state']
kickstarter_grading_df = kickstarter_grading_df.drop('state', axis=1)

categorical_features = kickstarter_grading_df.select_dtypes(include=['object']).columns
X1 = pd.get_dummies(kickstarter_grading_df, columns=categorical_features, drop_first=True)

# Find columns in X that are not in X1
columns_in_X_not_in_X1 = set(X.columns) - set(X1.columns)

# Append missing columns to X1 with default value of 0
for column in columns_in_X_not_in_X1:
    X1[column] = 0
    
# Ensure that the columns in X1 are in the same order as in X
X1 = X1[X.columns]

# Apply the model previously trained to the grading data
y_grading_pred = best_rf.predict(X1)

# Calculate the accuracy score
accuracy_score(y_grading, y_grading_pred)

#####################################################################
## CLUSTERING - FINAL MODEL
#####################################################################
import pandas as pd
import numpy as np
from sklearn.preprocessing import MinMaxScaler
from sklearn.metrics import calinski_harabasz_score 
from sklearn.cluster import KMeans,DBSCAN 
from sklearn.metrics import silhouette_score 
from matplotlib import pyplot as plt
from sklearn.ensemble import IsolationForest
from sklearn.decomposition import PCA
import seaborn as sns
# Import data
kickstarter_df = pd.read_excel("C:/Users/vinay/OneDrive/Documents/.spyder-py3/Kickstarter.xlsx")

# Pre-Processing
kickstarter_df = kickstarter_df.dropna()


# Filtering the DataFrame to keep only rows where 'status' is 'successful' or 'failed'
kickstarter_df = kickstarter_df[kickstarter_df['state'].isin(['successful', 'failed'])]

kickstarter_df['goal_fixed']=kickstarter_df['goal']*kickstarter_df['static_usd_rate']
kickstarter_df = kickstarter_df.drop(columns=['id', 'name','state_changed_at','pledged',
                                               'deadline', 'created_at', 'launched_at','goal','static_usd_rate',
                                               'deadline_weekday','state_changed_at_weekday','created_at_weekday','launched_at_weekday','deadline_month','deadline_day','deadline_yr','deadline_hr','state_changed_at_month','state_changed_at_day','state_changed_at_yr','state_changed_at_hr','created_at_month','created_at_day','created_at_yr','created_at_hr','launched_at_month','launched_at_day','launched_at_yr',
                                               'launched_at_hr','name_len','blurb_len'])
# Pre-Processing: Handling categorical variables
categorical_features = kickstarter_df.select_dtypes(include=['object']).columns
X = pd.get_dummies(kickstarter_df, columns=categorical_features, drop_first=True)

# Using Isolation Forest for detecting outliers
iso_forest = IsolationForest(n_estimators=100, contamination='auto', random_state=5)  # Let the algorithm determine the contamination
iso_forest.fit(X)

# Predict if a data point is an outlier. -1 for outliers, 1 for inliers.
outlier_predictions = iso_forest.predict(X)

# Select all rows that are not outliers
is_inlier = outlier_predictions == 1
X = X[is_inlier]

# Standardize the variables
# Initialize the MinMaxScaler
scaler = MinMaxScaler()

# Fit the scaler to your data and transform it
X_std = scaler.fit_transform(X)

# Convert the scaled data back to a DataFrame and assign column names
X_std = pd.DataFrame(X_std, columns=X.columns)


# Calculate inertia for each value of k
withinss = []
for i in range(2, 12):    
    kmeans = KMeans(n_clusters=i,random_state=5)
    model = kmeans.fit(X_std)
    withinss.append(model.inertia_)

# Create a plot
plt.plot(range(2, 12), withinss)
plt.title('Elbow Method')
plt.xlabel('Number of Clusters')
plt.ylabel('Within-Cluster Sum of Squares (WCSS)')
plt.show()

# Finding optimal K
for i in range (2,12):    
    kmeans = KMeans(n_clusters=i,random_state=5)
    model = kmeans.fit(X_std)
    labels = model.labels_
    print(i,':',silhouette_score(X_std,labels))
 
 # Calculate F-score

# Calculate Calinski-Harabasz Score
for i in range(2, 12):
    kmeans = KMeans(n_clusters=i,random_state=5)
    model = kmeans.fit(X_std)
    labels = model.labels_
    calinski_harabasz_avg = calinski_harabasz_score(X_std, labels)
    print(f'{i} Clusters: Calinski-Harabasz Score = {calinski_harabasz_avg}')


#4 is the optimal number of clusters
kmeans = KMeans(n_clusters=4,random_state=5)
model = kmeans.fit(X_std)
labels = model.labels_
centroids = model.cluster_centers_

# Convert centroids to a DataFrame for easier interpretation
centroids_df = pd.DataFrame(scaler.inverse_transform(centroids), columns=X_std.columns)

print("Cluster Centroids:")
print(centroids_df)

# Exporting centroids to a CSV file
centroids_df.to_csv("C:/Users/vinay/OneDrive/Documents/.spyder-py3/cluster_centroids.csv", index=False)
# Calculate the mean of all features across the entire dataset
overall_means = X_std.mean()

# Repeat the overall_means to match the number of clusters
overall_means_aligned = np.tile(overall_means.values, (centroids.shape[0], 1))

# Calculate how much each cluster's centroid deviates from the overall mean
feature_importance = pd.DataFrame(centroids - overall_means_aligned, columns=X_std.columns)

print("Feature Importance for Each Cluster:")
print(feature_importance)
feature_importance.to_csv("C:/Users/vinay/OneDrive/Documents/.spyder-py3/cluster_importance.csv", index=False)


##########################################################################################################
##### END OF CODE #######################################################################################
#########################################################################################################
##APPENDIX###
### EDA OF KICKSTARTER DATASET #########
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import pandas as pd

# Create a new Word document
doc = Document()
kickstarter_df = pd.read_excel("C:/Users/vinay/OneDrive/Documents/.spyder-py3/Kickstarter.xlsx")

for column in kickstarter_df.columns:
    # Skip non-numeric columns for histograms and boxplots
    if kickstarter_df[column].dtype in ['int64', 'float64']:
        plt.figure(figsize=(10, 6))
        sns.histplot(kickstarter_df[column], kde=True)
        plt.title(f'Distribution of {column}')
        plt.xlabel(column)
        plt.ylabel('Frequency')
        plot_filename = f'{column}_histogram.png'
        plt.savefig(plot_filename)
        doc.add_picture(plot_filename)
        plt.close()

        plt.figure(figsize=(10, 6))
        sns.boxplot(x=kickstarter_df[column])
        plt.title(f'Boxplot of {column}')
        plot_filename = f'{column}_boxplot.png'
        plt.savefig(plot_filename)
        doc.add_picture(plot_filename)
        plt.close()

    # For categorical data, use count plots
    elif kickstarter_df[column].dtype == 'object':
        plt.figure(figsize=(10, 6))
        sns.countplot(y=kickstarter_df[column])
        plt.title(f'Count Plot of {column}')
        plt.xlabel('Count')
        plt.ylabel(column)
        plot_filename = f'{column}_countplot.png'
        plt.savefig(plot_filename)
        doc.add_picture(plot_filename)
        plt.close()

# Save the document
doc_filename = 'kickstarter_eda_report.docx'
doc.save(doc_filename)

######## CLASSIFICATION MODEL SELECTION #####################
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler, OneHotEncoder
import pandas as pd
from sklearn.metrics import accuracy_score
from sklearn.neural_network import MLPClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.neighbors import KNeighborsClassifier

# Load data
kickstarter_df = pd.read_excel("C:/Users/vinay/OneDrive/Documents/.spyder-py3/Kickstarter.xlsx")

# Pre-Processing
kickstarter_df = kickstarter_df.dropna()
# Drop Future Information Columns from DataFrame  
kickstarter_df=kickstarter_df.drop(columns=['id','name','pledged','state_changed_at','backers_count','usd_pledged'
                                         ,'state_changed_at_weekday','state_changed_at_month','state_changed_at_day','state_changed_at_yr','state_changed_at_hr','launch_to_state_change_days'
                                         ,'deadline','created_at','launched_at'])
y=kickstarter_df['state']
kickstarter_df=kickstarter_df.drop('state', axis=1)

# Pre-Processing: Handling categorical variables
categorical_features = kickstarter_df.select_dtypes(include=['object']).columns
numerical_features = kickstarter_df.select_dtypes(exclude=['object']).columns

# Use pd.get_dummies to one-hot encode categorical variables
X = pd.get_dummies(kickstarter_df, columns=categorical_features, drop_first=True)
X.columns

# Splitting the dataset into the Training set and Test set
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=5)

# Standardizing features
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# Models
models = {
    "ANN": MLPClassifier(random_state=5),
    "Logistic Regression": LogisticRegression(random_state=5),
    "Random Forest": RandomForestClassifier(random_state=5),
    "Gradient Boosting": GradientBoostingClassifier(random_state=5),
    "KNN": KNeighborsClassifier(random_state=5)
}

# Training and Evaluating models
results = {}
for name, model in models.items():
    model.fit(X_train_scaled, y_train)
    predictions = model.predict(X_test_scaled)
    accuracy = accuracy_score(y_test, predictions)
    results[name] = accuracy

# Comparing Models
best_model = max(results, key=results.get)
print("Model Accuracies:", results)
print("Best Model:", best_model) 

